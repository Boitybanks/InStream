import io
import zipfile

from instream_ingestion import extract_attachment


def test_docx_extraction(tmp_path):
    import docx

    document = docx.Document()
    document.add_paragraph("Hello from a docx file")
    buffer = io.BytesIO()
    document.save(buffer)

    result = extract_attachment("statement.docx", buffer.getvalue())
    assert "Hello from a docx file" in result.text
    assert result.needs_ocr is False


def test_image_extraction_flags_needs_ocr():
    result = extract_attachment("scan.png", b"not-a-real-png-but-fine-for-this-test")
    assert result.needs_ocr is True
    assert result.text == ""


def test_zip_extraction_recurses_into_members():
    import docx

    document = docx.Document()
    document.add_paragraph("Nested document text")
    inner_buffer = io.BytesIO()
    document.save(inner_buffer)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as archive:
        archive.writestr("inner.docx", inner_buffer.getvalue())

    result = extract_attachment("bundle.zip", zip_buffer.getvalue())
    assert "Nested document text" in result.text


def test_unknown_extension_returns_empty():
    result = extract_attachment("data.xyz", b"whatever")
    assert result.text == ""
    assert result.needs_ocr is False
